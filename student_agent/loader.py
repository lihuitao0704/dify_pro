"""
文档加载器：解析 .txt / .docx / .xlsx / .pdf 多种格式
输出统一的知识库切片格式 [{title, content, source}]
"""

import os
import re


def load_txt(filepath: str) -> str:
    """读取纯文本"""
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()


def load_docx(filepath: str) -> str:
    """读取 Word 文档"""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # 也读表格
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                paragraphs.append(" | ".join(cells))
        return "\n\n".join(paragraphs)
    except ImportError:
        print(f"[Loader] python-docx 未安装，无法读取: {filepath}")
        return ""


def load_xlsx(filepath: str) -> str:
    """读取 Excel（每行一条记录）"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(filepath)
        lines = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                line = " | ".join(cells)
                if line.strip():
                    lines.append(line)
        return "\n".join(lines)
    except ImportError:
        print(f"[Loader] openpyxl 未安装，无法读取: {filepath}")
        return ""


def load_pdf(filepath: str) -> str:
    """读取 PDF"""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(filepath)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except ImportError:
        try:
            import pdfplumber
            text = ""
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text += t + "\n"
            return text
        except ImportError:
            print(f"[Loader] PyMuPDF/pdfplumber 未安装，无法读取: {filepath}")
            return ""


def load_file(filepath: str) -> str:
    """根据扩展名自动选择解析器"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in (".txt", ".md"):
        return load_txt(filepath)
    elif ext == ".docx":
        return load_docx(filepath)
    elif ext == ".xlsx":
        return load_xlsx(filepath)
    elif ext == ".pdf":
        return load_pdf(filepath)
    else:
        print(f"[Loader] 不支持的文件格式: {ext}")
        return ""


def load_folder(folder: str) -> list[dict]:
    """加载文件夹中所有支持的文件，返回 [{title, content, source}]"""
    docs = []
    if not os.path.isdir(folder):
        return docs

    for fname in os.listdir(folder):
        filepath = os.path.join(folder, fname)
        # 跳过临时文件
        if fname.startswith("~$"):
            continue
        content = load_file(filepath)
        if content:
            docs.append({
                "title": fname,
                "content": content,
                "source": filepath,
            })
            print(f"[Loader] 已加载: {fname} ({len(content)} 字符)")

    return docs


def split_chunks(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """按段落+长度切片"""
    paragraphs = text.split("\n\n")
    chunks = []
    current = ""
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current) + len(para) < chunk_size:
            current += para + "\n"
        else:
            if current:
                chunks.append(current.strip())
            current = para + "\n"
    if current:
        chunks.append(current.strip())
    return chunks


def extract_keywords(text: str) -> set[str]:
    """提取关键词（中文2-4字词组 + 英文3字母以上词）"""
    words = set()
    chinese_words = re.findall(r"[一-鿿]{2,4}", text)
    words.update(chinese_words)
    english_words = re.findall(r"[a-zA-Z]{3,}", text)
    words.update([w.lower() for w in english_words])
    return words
